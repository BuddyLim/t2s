"""Parse a Word (.docx) data dictionary into plain text for LLM context.

The data dictionary explains what each column/field means. Word documents vary
wildly in structure, so we extract both paragraph text and any tables (rendered
as simple "cell | cell" rows) and hand the whole thing to the model as context.
This is deliberately format-agnostic; it can be tightened once the real document
structure is known.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument


def _iter_block_text(doc: DocxDocument) -> list[str]:
    """Yield paragraph text and flattened table rows from the document."""
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    return lines


def load_dictionary(dict_path: str | Path) -> str:
    """Return the data dictionary as a single plain-text string.

    Raises FileNotFoundError if the document does not exist.
    """
    path = Path(dict_path)
    if not path.exists():
        raise FileNotFoundError(f"Data dictionary not found: {path}")

    doc = Document(str(path))
    return "\n".join(_iter_block_text(doc)).strip()
