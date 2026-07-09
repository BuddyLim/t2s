"""Tests for the Word (.docx) data dictionary parser."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from data_dictionary import load_dictionary

# Referencing a fixture by its name as a test argument is the pytest idiom.
# pylint: disable=redefined-outer-name


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a small .docx fixture with a paragraph and a table."""
    path = tmp_path / "dict.docx"
    doc = Document()
    doc.add_paragraph("Customer data dictionary")

    table = doc.add_table(rows=0, cols=2)
    for column, meaning in [
        ("Column", "Meaning"),
        ("name", "Customer full name"),
        ("age", "Age in years"),
    ]:
        cells = table.add_row().cells
        cells[0].text = column
        cells[1].text = meaning

    doc.save(str(path))
    return path


def test_extracts_paragraphs_and_tables(sample_docx: Path) -> None:
    """Paragraphs and table rows are both flattened into the dictionary text."""
    text = load_dictionary(sample_docx)
    assert "Customer data dictionary" in text
    assert "name | Customer full name" in text
    assert "age | Age in years" in text


def test_missing_file_raises(tmp_path: Path) -> None:
    """A missing .docx path raises FileNotFoundError."""
    with pytest.raises(FileNotFoundError):
        load_dictionary(tmp_path / "nope.docx")
